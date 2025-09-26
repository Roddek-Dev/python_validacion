"""
Extractor de contenido de múltiples formatos de archivo.
"""

import logging
from pathlib import Path
from typing import Dict, Optional

logger = logging.getLogger(__name__)


class ContentExtractor:
    """Extrae texto de diferentes formatos de archivo."""
    
    def __init__(self, capabilities: Dict[str, bool], config: Dict):
        self.capabilities = capabilities
        self.config = config
        self.extractors = self._initialize_extractors()
    
    def _initialize_extractors(self) -> Dict[str, callable]:
        """Inicializa extractores disponibles basados en capabilities."""
        extractors = {}
        
        if self.capabilities.get('pdf_processing'):
            extractors['.pdf'] = self._extract_from_pdf
        
        if self.capabilities.get('docx_processing'):
            extractors['.docx'] = self._extract_from_docx
        
        if self.capabilities.get('xlsx_processing'):
            extractors['.xlsx'] = self._extract_from_xlsx
        
        # TXT siempre disponible
        extractors['.txt'] = self._extract_from_txt
        
        return extractors
    
    def extract_text(self, file_path: Path) -> str:
        """
        Extrae texto de un archivo según su extensión.
        
        Args:
            file_path: Ruta al archivo
            
        Returns:
            Texto extraído del archivo
        """
        extension = file_path.suffix.lower()
        
        if extension not in self.extractors:
            logger.debug(f"No hay extractor disponible para: {extension}")
            return ""
        
        try:
            return self.extractors[extension](file_path)
        except Exception as e:
            logger.error(f"Error extrayendo contenido de {file_path}: {e}")
            return ""
    
    def can_extract(self, file_path: Path) -> bool:
        """
        Verifica si se puede extraer contenido del archivo.
        
        Args:
            file_path: Ruta al archivo
            
        Returns:
            True si se puede extraer contenido
        """
        extension = file_path.suffix.lower()
        return extension in self.extractors
    
    def _extract_from_pdf(self, file_path: Path) -> str:
        """Extrae texto de archivo PDF."""
        try:
            import pdfplumber
            
            text = ""
            max_pages = self.config.get("pdf_pages_to_read", 1)
            
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages[:max_pages]):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"

                        # Extraer texto de tablas
                        tables = page.extract_tables()
                        for table in tables:
                            for row in table:
                                if row:
                                    text += " ".join(str(cell) for cell in row if cell) + "\n"
                    except Exception as page_error:
                        logger.warning(f"Error procesando página {i+1} de {file_path}: {page_error}")
                        continue
            
            return text
            
        except Exception as e:
            logger.error(f"Error extrayendo texto de PDF {file_path}: {e}")
            return ""
    
    def _extract_from_docx(self, file_path: Path) -> str:
        """Extrae texto de archivo DOCX."""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extraer texto de tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
            
        except Exception as e:
            logger.error(f"Error extrayendo texto de DOCX {file_path}: {e}")
            return ""
    
    def _extract_from_xlsx(self, file_path: Path) -> str:
        """Extrae texto de archivo XLSX."""
        try:
            from openpyxl import load_workbook
            
            workbook = load_workbook(file_path, data_only=True)
            text = ""
            max_sheets = self.config.get("xlsx_max_sheets", 3)
            
            for i, sheet_name in enumerate(workbook.sheetnames[:max_sheets]):
                sheet = workbook[sheet_name]
                
                # Limitar número de filas para archivos grandes
                max_rows = min(sheet.max_row, 1000)
                
                for row in sheet.iter_rows(max_row=max_rows, values_only=True):
                    for cell in row:
                        if cell is not None:
                            text += str(cell) + " "
                text += "\n"
            
            return text
            
        except Exception as e:
            logger.error(f"Error extrayendo texto de XLSX {file_path}: {e}")
            return ""
    
    def _extract_from_txt(self, file_path: Path) -> str:
        """Extrae texto de archivo TXT."""
        try:
            # Intentar diferentes encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, "r", encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            
            # Si ningún encoding funciona, leer con errores ignorados
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
                
        except Exception as e:
            logger.error(f"Error extrayendo texto de TXT {file_path}: {e}")
            return ""
    
    def get_supported_extensions(self) -> list:
        """Retorna lista de extensiones soportadas."""
        return list(self.extractors.keys())
    
    def get_extraction_stats(self) -> Dict[str, int]:
        """Retorna estadísticas de extracción (para implementar en el futuro)."""
        return {
            "total_extractions": 0,
            "successful_extractions": 0,
            "failed_extractions": 0,
            "avg_extraction_time": 0.0
        }